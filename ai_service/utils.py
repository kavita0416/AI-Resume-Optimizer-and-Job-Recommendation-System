# ai_service/utils.py
from PyPDF2 import PdfReader
import io
import docx
import re

def extract_text_from_bytes(file_bytes, filename):
    """
    Detects file type and extracts text accordingly (PDF, DOCX, DOC, TXT).
    Returns clean plain text.
    """
    filename = filename.lower()

    try:
        if filename.endswith(".pdf"):
            text = extract_text_pdf(file_bytes)
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            text = extract_text_docx(file_bytes)
        elif filename.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            # fallback try utf-8 decode
            text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"❌ Error reading file ({filename}):", e)
        text = ""

    return clean_text(text)


def extract_text_pdf(file_bytes):
    """
    Extracts text from PDF file bytes using PyPDF2.
    Returns joined text from all pages.
    """
    text = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            except Exception as e:
                print("⚠️ Error extracting page text:", e)
    except Exception as e:
        print("❌ Failed to read PDF:", e)

    joined_text = "\n".join(text)
    if not joined_text.strip():
        # Optional: placeholder for future OCR integration
        print("⚠️ No extractable text found (possibly scanned PDF). Consider OCR fallback.")
    return joined_text


def extract_text_docx(file_bytes):
    """
    Extracts text from DOCX file bytes using python-docx.
    Includes text from tables and paragraphs.
    """
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        print("❌ Error opening DOCX file:", e)
        return ""

    paragraphs = []
    # Normal paragraphs
    for p in document.paragraphs:
        if p.text and p.text.strip():
            paragraphs.append(p.text.strip())

    # Table content
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def clean_text(text: str) -> str:
    """
    Cleans and normalizes extracted text:
    - Removes extra spaces
    - Collapses multiple newlines
    - Strips non-printable chars
    """
    if not text:
        return ""

    # Remove control chars and normalize whitespace
    text = re.sub(r"[\x00-\x1F\x7F]", " ", text)
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    # Trim each line
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    cleaned = "\n".join(lines).strip()

    return cleaned
